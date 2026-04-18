from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 본문 기본 색(인쇄·뷰어에서 자동색으로 흐려지는 현상 방지)
_BODY_RGB = RGBColor(0x1A, 0x1A, 0x1A)

# ── 기본 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
font.color.rgb = _BODY_RGB
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 페이지 설정
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text, bold=False, size=10, color=None, space_after=6, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color if color is not None else _BODY_RGB
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    if alignment:
        p.alignment = alignment
    return p


def add_bullet(doc, text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    if p.runs:
        p.runs[0].text = text
        run = p.runs[0]
        for r in p.runs[1:]:
            r.text = ''
    else:
        run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.color.rgb = _BODY_RGB
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    return p


def add_table_with_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 헤더
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A1A2E')

    # 데이터
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '맑은 고딕'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    run.font.color.rgb = _BODY_RGB
                paragraph.paragraph_format.space_after = Pt(2)
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F4F8')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para(doc, "제안요청서", bold=True, size=28, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc, "(RFP: Request for Proposal)", bold=False, size=12, color=RGBColor(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para(doc, "═" * 50, size=10, color=RGBColor(0xCC, 0xCC, 0xCC), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_para(doc, "프레시밀 브랜드 리뉴얼 및 D2C 이커머스 플랫폼 구축", bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para(doc, "발주처: (주)프레시밀 마케팅본부", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "문서번호: FM-RFP-2025-003", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "작성일: 2025년 4월 15일", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "보안등급: 대외비", size=11, color=RGBColor(0xCC, 0x33, 0x33), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

doc.add_page_break()


# ══════════════════════════════════════════════════
# 목차
# ══════════════════════════════════════════════════
add_heading_styled(doc, "목차", level=1)
add_para(doc, "", space_after=6)

toc_items = [
    "1. 사업 개요",
    "    1.1 사업 배경 및 목적",
    "    1.2 사업 범위",
    "    1.3 기대 효과",
    "2. 현황 분석",
    "    2.1 기존 시스템 현황",
    "    2.2 주요 문제점",
    "3. 요구사항 정의",
    "    3.1 브랜드 리뉴얼 요구사항",
    "    3.2 D2C 이커머스 플랫폼 요구사항",
    "    3.3 통합 관리 시스템 요구사항",
    "    3.4 마케팅·데이터 연동 요구사항",
    "4. 기술 요건",
    "5. 사업 추진 일정",
    "6. 제안서 작성 지침",
    "7. 평가 기준",
    "8. 제안 참가 자격",
    "9. 기타 사항",
    "    9.1 예산",
    "    9.2 계약 조건",
    "    9.3 보안 및 기밀유지",
    "    9.4 문의처",
]

for item in toc_items:
    if item.startswith("    "):
        add_para(doc, item, size=10, color=RGBColor(0x44, 0x44, 0x44), space_after=3)
    else:
        add_para(doc, item, size=10, bold=True, space_after=3)

doc.add_page_break()


# ══════════════════════════════════════════════════
# 1. 사업 개요
# ══════════════════════════════════════════════════
add_heading_styled(doc, "1. 사업 개요", level=1)

add_heading_styled(doc, "1.1 사업 배경 및 목적", level=2)

add_para(doc, "(주)프레시밀은 2018년에 설립된 프리미엄 밀키트·간편식 브랜드로, 현재 네이버 스마트스토어와 쿠팡을 중심으로 연 매출 약 85억 원 규모의 사업을 운영하고 있다. 최근 3년간 연평균 40% 이상의 매출 성장을 기록하고 있으나, 다음과 같은 구조적 한계에 직면해 있다.")

add_bullet(doc, "마켓플레이스 수수료 부담 증가 (평균 15~22%)로 수익성 악화")
add_bullet(doc, "고객 데이터 확보 한계로 인한 재구매율 정체 (현재 18.5%)")
add_bullet(doc, "2020년 런칭 당시의 브랜드 아이덴티티가 현재 시장 포지셔닝과 괴리가 있다")
add_bullet(doc, "경쟁사 대비 디지털 고객 경험(CX) 열위")

add_para(doc, "이에 브랜드 리뉴얼을 통해 프리미엄 포지셔닝을 강화하고, 자사 D2C(Direct to Consumer) 이커머스 플랫폼을 구축하여 고객과의 직접적인 관계를 확보하고자 한다.")

add_para(doc, "본 사업의 핵심 목적은 다음과 같다.")
add_bullet(doc, "브랜드 아이덴티티 재정립 및 전 채널 일관된 브랜드 경험 구현")
add_bullet(doc, "자사 D2C 이커머스 플랫폼 구축을 통한 마켓플레이스 의존도 탈피")
add_bullet(doc, "고객 데이터 기반의 개인화 마케팅 체계 구축")
add_bullet(doc, "구독 모델 도입을 통한 안정적 반복 매출 구조 확보")

add_heading_styled(doc, "1.2 사업 범위", level=2)

add_para(doc, "본 사업의 범위는 크게 4개 영역으로 구분된다.")

add_table_with_data(doc,
    ["구분", "영역", "주요 내용"],
    [
        ["1", "브랜드 리뉴얼", "BI/CI 리뉴얼, 브랜드 가이드라인, 패키지 디자인 리뉴얼"],
        ["2", "D2C 이커머스 플랫폼", "반응형 웹사이트, 회원/주문/결제/배송, 구독 시스템"],
        ["3", "통합 관리 시스템", "어드민, 상품/재고 관리, CS 관리, 정산"],
        ["4", "마케팅·데이터 연동", "GA4, CRM, 마케팅 자동화, 리뷰 시스템"],
    ],
    col_widths=[1.5, 4, 10.5]
)

add_heading_styled(doc, "1.3 기대 효과", level=2)

add_table_with_data(doc,
    ["항목", "현재", "목표 (런칭 12개월 후)"],
    [
        ["D2C 매출 비중", "0%", "전체 매출의 25% 이상"],
        ["고객 재구매율", "18.5%", "35% 이상"],
        ["마켓플레이스 수수료율", "평균 18%", "실질 수수료 8% 이하 (PG 수수료)"],
        ["고객 데이터 확보율", "약 12% (리뷰 작성 기준)", "회원 가입 기준 70% 이상"],
        ["구독 고객 수", "0명", "3,000명 이상"],
    ],
    col_widths=[4.5, 5, 6.5]
)


# ══════════════════════════════════════════════════
# 2. 현황 분석
# ══════════════════════════════════════════════════
add_heading_styled(doc, "2. 현황 분석", level=1)

add_heading_styled(doc, "2.1 기존 시스템 현황", level=2)

add_table_with_data(doc,
    ["항목", "현황", "비고"],
    [
        ["주요 판매채널", "네이버 스마트스토어, 쿠팡, SSG", "자사몰 없음"],
        ["브랜드 웹사이트", "WordPress 기반 소개 페이지", "결제 기능 없음, 모바일 미최적화"],
        ["상품 관리", "채널별 개별 관리 (수동)", "사방넷 일부 연동"],
        ["고객 관리", "엑셀 기반 수동 관리", "CRM 시스템 없음"],
        ["마케팅", "네이버 검색광고, 메타 광고 집행", "성과 측정 통합 대시보드 없음"],
        ["디자인 자산", "로고, 패키지 (2020년 제작)", "브랜드 가이드라인 미비"],
    ],
    col_widths=[3.5, 6, 6.5]
)

add_heading_styled(doc, "2.2 주요 문제점", level=2)

add_bullet(doc, "채널 분산으로 인한 통합 고객 경험 부재: 각 마켓플레이스마다 상이한 상세페이지, 프로모션, 고객 응대 기준으로 브랜드 일관성 저하")
add_bullet(doc, "고객 데이터 사일로: 채널별로 고객 데이터가 분리되어 있어 통합 분석 및 개인화 마케팅 불가")
add_bullet(doc, "브랜드 자산 노후화: 현재 BI가 '가정간편식' 이미지에 머물러 있어 '프리미엄 푸드 라이프스타일' 포지셔닝과 불일치한다")
add_bullet(doc, "운영 비효율: 채널별 수동 상품 등록, 재고 관리, CS 대응으로 운영 인력 과부하")


# ══════════════════════════════════════════════════
# 3. 요구사항 정의
# ══════════════════════════════════════════════════
add_heading_styled(doc, "3. 요구사항 정의", level=1)

add_heading_styled(doc, "3.1 브랜드 리뉴얼 요구사항", level=2)

add_para(doc, "[BR-001] 브랜드 아이덴티티 재수립", bold=True, size=10, space_after=3)
add_bullet(doc, "브랜드 미션, 비전, 핵심 가치 재정립")
add_bullet(doc, "브랜드 페르소나 및 톤앤매너 정의")
add_bullet(doc, "브랜드 스토리 개발")
add_bullet(doc, "경쟁사 분석 및 차별화 포지셔닝 전략 수립")

add_para(doc, "[BR-002] 비주얼 아이덴티티 개발", bold=True, size=10, space_after=3)
add_bullet(doc, "로고 리디자인 (워드마크 + 심볼 + 시그니처 조합)")
add_bullet(doc, "컬러 시스템 (Primary, Secondary, Accent)")
add_bullet(doc, "타이포그래피 시스템 (국문, 영문)")
add_bullet(doc, "그래픽 모티프 및 일러스트레이션 스타일")
add_bullet(doc, "포토그래피 가이드라인")

add_para(doc, "[BR-003] 브랜드 가이드라인 문서", bold=True, size=10, space_after=3)
add_bullet(doc, "브랜드 가이드라인 문서 (최소 60페이지 이상)")
add_bullet(doc, "디지털 적용 가이드 (웹, 앱, SNS)")
add_bullet(doc, "패키지 디자인 적용 가이드")
add_bullet(doc, "오프라인 적용 가이드 (명함, 리플렛, 전시 부스 등)")

add_para(doc, "[BR-004] 패키지 디자인 리뉴얼", bold=True, size=10, space_after=3)
add_bullet(doc, "주력 상품 라인 3개(밀키트, 간편식, 소스류) 패키지 리디자인")
add_bullet(doc, "각 라인별 대표 상품 2종, 총 6종 패키지 디자인")
add_bullet(doc, "택배 박스, 보냉백, 동봉 인쇄물 디자인")

add_heading_styled(doc, "3.2 D2C 이커머스 플랫폼 요구사항", level=2)

add_para(doc, "[EC-001] 프론트엔드 (고객 대면)", bold=True, size=10, space_after=3)
add_bullet(doc, "반응형 웹 (PC, 태블릿, 모바일 완벽 대응)")
add_bullet(doc, "메인 페이지: 히어로 배너, 베스트셀러, 신상품, 브랜드 스토리 섹션")
add_bullet(doc, "상품 목록: 카테고리 필터, 정렬, 태그 기반 탐색")
add_bullet(doc, "상품 상세: 이미지 갤러리, 옵션 선택, 영양 정보, 조리법, 리뷰")
add_bullet(doc, "장바구니 및 위시리스트")
add_bullet(doc, "매거진/레시피 콘텐츠 섹션")
add_bullet(doc, "검색 기능 (자동완성, 연관 상품 추천)")

add_para(doc, "[EC-002] 회원 시스템", bold=True, size=10, space_after=3)
add_bullet(doc, "회원가입 (이메일, 소셜 로그인: 카카오, 네이버, 구글, 애플)")
add_bullet(doc, "마이페이지: 주문내역, 배송추적, 구독 관리, 포인트, 쿠폰, 1:1 문의")
add_bullet(doc, "등급제: 4단계 (일반 → 실버 → 골드 → VIP), 등급별 혜택 차등")
add_bullet(doc, "포인트 시스템: 구매 적립, 리뷰 적립, 출석 체크, 친구 추천")

add_para(doc, "[EC-003] 주문 및 결제", bold=True, size=10, space_after=3)
add_bullet(doc, "PG 연동: 토스페이먼츠 (신용카드, 계좌이체, 가상계좌, 간편결제)")
add_bullet(doc, "간편결제: 카카오페이, 네이버페이, 토스페이")
add_bullet(doc, "주문 상태 관리: 주문접수 → 결제완료 → 상품준비 → 배송중 → 배송완료")
add_bullet(doc, "부분 취소/반품/교환 처리")
add_bullet(doc, "선물하기 기능 (수령인 정보 별도 입력)")

add_para(doc, "[EC-004] 구독 시스템 ★ 핵심 기능", bold=True, size=10, space_after=3)
add_bullet(doc, "구독 주기 선택: 1주/2주/3주/4주")
add_bullet(doc, "구독 상품 큐레이션: 고객 취향 설문 기반 자동 추천")
add_bullet(doc, "구독 관리: 건너뛰기, 일시정지, 상품 변경, 주기 변경, 해지")
add_bullet(doc, "구독 전용 할인 (정기구독 시 10~15% 할인)")
add_bullet(doc, "구독 고객 전용 콘텐츠 및 사전 출시 상품 접근")
add_bullet(doc, "정기결제 자동 처리 및 결제 실패 시 재시도/알림 로직")

add_para(doc, "[EC-005] 배송", bold=True, size=10, space_after=3)
add_bullet(doc, "배송 추적 연동 (CJ대한통운, 한진택배)")
add_bullet(doc, "새벽배송/일반배송 선택 (수도권 새벽배송)")
add_bullet(doc, "배송일 지정 기능")
add_bullet(doc, "배송비 정책: 3만원 이상 무료배송, 제주/도서산간 추가요금")
add_bullet(doc, "배송 완료 후 자동 리뷰 요청 알림")

add_heading_styled(doc, "3.3 통합 관리 시스템 요구사항", level=2)

add_para(doc, "[AD-001] 관리자 대시보드", bold=True, size=10, space_after=3)
add_bullet(doc, "실시간 매출/주문/방문자 현황 대시보드")
add_bullet(doc, "일별/주별/월별 매출 추이, 상품별 판매 현황")
add_bullet(doc, "구독 관련 지표: 신규 구독, 해지율, 구독 매출 비중")
add_bullet(doc, "고객 유입 채널별 분석")

add_para(doc, "[AD-002] 상품 관리", bold=True, size=10, space_after=3)
add_bullet(doc, "상품 등록/수정/삭제 (옵션, 가격, 이미지, 상세 설명)")
add_bullet(doc, "카테고리 및 태그 관리")
add_bullet(doc, "재고 관리 및 품절 자동 처리")
add_bullet(doc, "상품 노출 순서 및 기획전 관리")

add_para(doc, "[AD-003] 주문/CS 관리", bold=True, size=10, space_after=3)
add_bullet(doc, "주문 목록 및 상태 관리 (검색, 필터, 일괄 처리)")
add_bullet(doc, "취소/반품/교환 처리 워크플로우")
add_bullet(doc, "1:1 문의 관리 및 답변 템플릿")
add_bullet(doc, "FAQ 관리")

add_heading_styled(doc, "3.4 마케팅·데이터 연동 요구사항", level=2)

add_para(doc, "[MK-001] 분석 및 추적", bold=True, size=10, space_after=3)
add_bullet(doc, "Google Analytics 4 (GA4) 전자상거래 추적 연동")
add_bullet(doc, "메타 픽셀, 카카오 픽셀, 네이버 전환 추적 스크립트 설치")
add_bullet(doc, "UTM 파라미터 기반 유입 채널 추적")

add_para(doc, "[MK-002] CRM 및 마케팅 자동화", bold=True, size=10, space_after=3)
add_bullet(doc, "카카오 알림톡/친구톡 연동 (주문 알림, 배송 알림, 마케팅 메시지)")
add_bullet(doc, "이메일 마케팅 연동 (스티비 또는 동급)")
add_bullet(doc, "자동 시나리오: 장바구니 이탈 알림, 재구매 유도, 구독 만료 알림, 생일 쿠폰")
add_bullet(doc, "고객 세그먼트 기반 타겟 마케팅")

add_para(doc, "[MK-003] 리뷰 시스템", bold=True, size=10, space_after=3)
add_bullet(doc, "텍스트 + 포토 리뷰 작성")
add_bullet(doc, "리뷰 작성 시 포인트 자동 적립")
add_bullet(doc, "리뷰 관리 (승인, 숨김, 베스트 리뷰 선정)")
add_bullet(doc, "리뷰 기반 상품 평점 자동 계산")


# ══════════════════════════════════════════════════
# 4. 기술 요건
# ══════════════════════════════════════════════════
add_heading_styled(doc, "4. 기술 요건", level=1)

add_table_with_data(doc,
    ["항목", "요건", "비고"],
    [
        ["프레임워크", "React 또는 Next.js 기반 프론트엔드\nNode.js 또는 Python 기반 백엔드", "제안사 자유 (근거 제시)"],
        ["호스팅", "AWS 또는 동급 클라우드", "Auto Scaling 필수"],
        ["데이터베이스", "RDBMS + NoSQL 혼합 구조 권장", ""],
        ["보안", "SSL 인증서, 개인정보 암호화,\nOWASP Top 10 대응", "개인정보보호법 준수 필수"],
        ["성능", "페이지 로딩 3초 이내,\n동시 접속 5,000명 처리", "Core Web Vitals 기준 충족"],
        ["CDN", "이미지/정적 파일 CDN 적용", "CloudFront 또는 동급"],
        ["모니터링", "서버 모니터링, 에러 트래킹, 로그 관리", ""],
        ["백업", "일 1회 자동 백업, 30일 보관", ""],
    ],
    col_widths=[3, 7.5, 5.5]
)


# ══════════════════════════════════════════════════
# 5. 사업 추진 일정
# ══════════════════════════════════════════════════
add_heading_styled(doc, "5. 사업 추진 일정", level=1)

add_para(doc, "전체 사업 기간: 계약일로부터 20주 (약 5개월)")
add_para(doc, "희망 착수일: 2025년 5월 19일")
add_para(doc, "희망 완료일: 2025년 10월 10일")

add_table_with_data(doc,
    ["단계", "기간", "주요 내용", "산출물"],
    [
        ["1단계\n킥오프 및 전략 수립", "1~3주", "프로젝트 킥오프, 브랜드 전략 워크숍,\n벤치마킹 및 경쟁사 분석, IA 설계", "전략 보고서, IA 문서"],
        ["2단계\n브랜드 리뉴얼", "3~8주", "브랜드 아이덴티티 개발,\n비주얼 시스템 디자인,\n패키지 디자인 (병행 진행)", "브랜드 가이드라인,\n패키지 시안"],
        ["3단계\nUX/UI 설계", "5~10주", "와이어프레임, UI 디자인,\n프로토타입, 사용성 테스트", "디자인 시안,\n프로토타입"],
        ["4단계\n개발", "8~17주", "프론트엔드/백엔드 개발,\n결제/배송/구독 시스템 구현,\n관리자 시스템 개발", "개발 산출물"],
        ["5단계\n테스트 및 오픈", "17~20주", "통합 테스트, 부하 테스트,\n보안 점검, 오픈 및 안정화", "테스트 보고서,\n최종 산출물"],
    ],
    col_widths=[3, 2, 5.5, 3.5]
)

add_para(doc, "※ 브랜드 리뉴얼과 UX/UI 설계는 일부 기간이 중첩되어 병행 진행된다.", size=9, color=RGBColor(0x55, 0x55, 0x55))
add_para(doc, "※ 각 단계별 중간 보고 및 산출물 검수를 진행하며, 상세 일정은 착수 후 협의한다.", size=9, color=RGBColor(0x55, 0x55, 0x55))


# ══════════════════════════════════════════════════
# 6. 제안서 작성 지침
# ══════════════════════════════════════════════════
add_heading_styled(doc, "6. 제안서 작성 지침", level=1)

add_para(doc, "제안서는 아래 목차에 따라 작성해 주시기 바랍니다.")

add_table_with_data(doc,
    ["순서", "항목", "주요 포함 내용", "권장 분량"],
    [
        ["1", "회사 소개", "회사 연혁, 조직 구성, 핵심 역량,\n관련 수상/인증 내역", "5페이지 이내"],
        ["2", "프로젝트 이해", "사업 배경에 대한 이해, RFP 분석,\n핵심 이슈 도출", "5~8페이지"],
        ["3", "브랜드 전략 제안", "브랜드 포지셔닝 방향, 리뉴얼 전략,\n무드보드 및 컨셉 키워드", "8~12페이지"],
        ["4", "UX/UI 전략 및 설계 방향", "정보 구조(IA), 핵심 화면 와이어프레임,\nUI 디자인 컨셉, 사용자 여정 맵", "10~15페이지"],
        ["5", "기술 제안", "시스템 아키텍처, 기술 스택 선정 근거,\n구독 시스템 구현 방안, 보안 방안", "8~12페이지"],
        ["6", "수행 방안", "프로젝트 추진 체계, 단계별 수행 방안,\n품질 관리, 커뮤니케이션 방안", "5~8페이지"],
        ["7", "투입 인력", "PM, 디자이너, 개발자 등\n투입 인력 프로필 및 역할", "3~5페이지"],
        ["8", "수행 일정", "WBS, 마일스톤, 주요 일정 계획", "2~3페이지"],
        ["9", "포트폴리오", "유사 프로젝트 수행 사례 3건 이상\n(브랜딩+이커머스 통합 사례 우대)", "5~8페이지"],
        ["10", "견적", "항목별 상세 견적서", "별도 첨부"],
    ],
    col_widths=[1.2, 3.5, 6, 3]
)

add_para(doc, "※ 전체 분량: 60~80페이지 이내 (A4 기준, 견적서 별도)")
add_para(doc, "※ 파일 형식: PPT (발표용) + PDF (검토용) 각 1부")
add_para(doc, "※ 제출 기한: 2025년 5월 9일 (금) 18:00까지")
add_para(doc, "※ 제출처: proposal@freshmeal.co.kr")


# ══════════════════════════════════════════════════
# 7. 평가 기준
# ══════════════════════════════════════════════════
add_heading_styled(doc, "7. 평가 기준", level=1)

add_table_with_data(doc,
    ["평가 영역", "평가 항목", "배점", "세부 기준"],
    [
        ["브랜드 전략\n(25점)", "브랜드 이해도", "10점", "자사 브랜드 및 시장에 대한 분석 깊이"],
        ["", "리뉴얼 방향성", "10점", "전략의 차별성과 실현 가능성"],
        ["", "크리에이티브 역량", "5점", "무드보드, 컨셉의 완성도"],
        ["기술 역량\n(25점)", "아키텍처 설계", "10점", "시스템 구조의 확장성, 안정성"],
        ["", "구독 시스템 설계", "10점", "구독 모델 구현 방안의 구체성"],
        ["", "기술 스택 적합성", "5점", "기술 선택의 근거와 타당성"],
        ["수행 능력\n(25점)", "유사 프로젝트 경험", "15점", "브랜딩+이커머스 통합 수행 경험"],
        ["", "투입 인력 역량", "10점", "핵심 인력의 전문성과 경험"],
        ["가격\n(15점)", "견적 적정성", "15점", "범위 대비 견적의 합리성"],
        ["프레젠테이션\n(10점)", "발표 및 Q&A", "10점", "전달력, 질의 대응 능력"],
    ],
    col_widths=[3, 3, 1.5, 7.5]
)

add_para(doc, "※ 총점 100점 만점이며, 기술 평가 85점과 가격 평가 15점으로 구성된다.")
add_para(doc, "※ 기술 평가 70점 미만인 업체는 가격 평가 대상에서 제외된다.")


# ══════════════════════════════════════════════════
# 8. 제안 참가 자격
# ══════════════════════════════════════════════════
add_heading_styled(doc, "8. 제안 참가 자격", level=1)

add_bullet(doc, "설립 후 3년 이상 경과한 법인 사업자")
add_bullet(doc, "최근 3년 내 브랜드 리뉴얼 및 이커머스 구축을 동시에 수행한 실적 1건 이상")
add_bullet(doc, "최근 3년 내 연 매출 50억 원 이상 브랜드의 이커머스 구축 실적 1건 이상")
add_bullet(doc, "상시 근무 인력 15인 이상 (디자인, 개발 인력 각 5인 이상)")
add_bullet(doc, "구독 커머스 또는 정기결제 시스템 구축 경험 보유 (우대)")
add_bullet(doc, "식품/F&B 업종 프로젝트 수행 경험 보유 (우대)")


# ══════════════════════════════════════════════════
# 9. 기타 사항
# ══════════════════════════════════════════════════
add_heading_styled(doc, "9. 기타 사항", level=1)

add_heading_styled(doc, "9.1 예산", level=2)
add_para(doc, "본 사업의 총 예산은 VAT 별도 2억 5천만 원 이내로 한다. 브랜드 리뉴얼과 이커머스 구축 비용을 항목별로 분리하여 제시하되, 통합 할인액이 있는 경우에는 별도로 표기하기 바란다.")

add_table_with_data(doc,
    ["구분", "예산 범위 (VAT 별도)"],
    [
        ["브랜드 리뉴얼 (BI/CI, 가이드라인, 패키지)", "7,000만 원 ~ 9,000만 원"],
        ["D2C 이커머스 플랫폼 (프론트/백엔드/관리자)", "1억 2,000만 원 ~ 1억 4,000만 원"],
        ["마케팅/데이터 연동", "2,000만 원 ~ 3,000만 원"],
        ["총 예산 한도", "2억 5,000만 원 이내"],
    ],
    col_widths=[8, 8]
)

add_heading_styled(doc, "9.2 계약 조건", level=2)
add_bullet(doc, "계약 방식: 총액 계약 (단계별 마일스톤 검수)")
add_bullet(doc, "대금 지급: 착수금 30%, 중간금 40% (3단계 완료 시), 잔금 30% (최종 검수 완료 시)")
add_bullet(doc, "하자 보수: 최종 검수 완료일로부터 12개월")
add_bullet(doc, "유지보수: 별도 유지보수 계약 협의 (월 정액 방식)")
add_bullet(doc, "저작권: 최종 납품된 결과물의 저작재산권은 (주)프레시밀에 귀속된다")

add_heading_styled(doc, "9.3 보안 및 기밀유지", level=2)
add_para(doc, "본 RFP에 포함된 모든 정보는 대외비이며, 제안 목적 외 사용을 금지한다. 제안에 참여하는 모든 업체는 제안 참여 전 비밀유지서약서(NDA)를 제출해야 한다. 선정되지 않은 업체에 제공된 자료는 선정 결과 통보 후 7일 이내에 폐기한다.")

add_heading_styled(doc, "9.4 문의처", level=2)

add_table_with_data(doc,
    ["항목", "내용"],
    [
        ["담당 부서", "마케팅본부 디지털전략팀"],
        ["담당자", "김서연 팀장"],
        ["이메일", "sy.kim@freshmeal.co.kr"],
        ["전화", "02-1234-5678 (내선 302)"],
        ["문의 기한", "2025년 5월 2일 (금) 18:00까지"],
        ["질의 방식", "이메일 접수 → 취합 후 일괄 답변 (5/7 예정)"],
    ],
    col_widths=[4, 12]
)

add_para(doc, "")
add_para(doc, "― 이상 ―", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)


# ── 저장 ──
filename = "샘플_RFP_프레시밀_브랜드리뉴얼_D2C구축.docx"
doc.save(filename)
print(f"✅ RFP 문서 생성 완료: {filename}")
print(f"📄 총 분량: 약 12페이지")
